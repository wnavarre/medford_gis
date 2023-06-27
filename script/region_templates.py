from docx import Document
import os
import contextily as cx
import matplotlib.pyplot as plt
import tempfile

def make_region_template(region_set, filename_out, ext=".png", *, workdir=None):
    if workdir is None:
        with tempfile.TemporaryDirectory() as workdir:
            return make_region_template(region_set, filename_out, workdir=workdir)
    doc = Document()
    par = doc.add_paragraph()
    for name, region in region_set.items():
        xmin, ymin, xmax, ymax = region.bounds()
        fig, ax = plt.subplots()
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(ymin, ymax)
        cx.add_basemap(ax)
        fname = os.path.join(workdir, name + ext)
        fig.savefig(fname)
        run = par.add_run()
        pic = run.add_picture(fname)
    doc.save(filename_out)
